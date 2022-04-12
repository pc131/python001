import random
from random import randint
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

working_dir = "C:\\Users\\tomasz.skoczylas\\Downloads\\11\\"
cw_working_dir = working_dir + "Creative_Work"

mylist = os.listdir(cw_working_dir)

period = '04_2022'
period_working_days = 21
period_working_hours = period_working_days * 8
number_of_works = len(mylist)

#list to keep hours for particular works
#divide working hours a little randomly
particular_hours = []

working_hours_even = period_working_hours % number_of_works

particular_works_hours  = int(period_working_hours / number_of_works)

working_hours_odd = particular_works_hours + (period_working_hours % number_of_works)

if working_hours_even != 0:
    for a in range(number_of_works - 1):
        particular_hours.append(particular_works_hours)
    particular_hours.append(working_hours_odd)
else:
    for a in range(number_of_works):
        particular_hours.append(particular_works_hours) 

#divide working hours a little randomly
all_list = []

def num_pieces(num,lenght):
    for i in range(lenght-1):
        if i == 0:
            n = random.randint(1,num)
            if n < 0.8*num/lenght or n> 1.2*num/lenght:
                n = int(0.9*num/lenght)
            all_list.append(n)
            num -= n
        else:
            n = random.randint(1,num)
            if n < 0.8*num/(lenght-1) or n> 1.2*num/(lenght-1):
                n = int(0.9*num/(lenght-1))
            all_list.append(n)
            num -= n
    all_list.append(num) 

num_pieces(int(0.9*period_working_hours), number_of_works) 
#divide working hours a little randomly

def w_miesiacu(month):
    match month:
        case '01':
            return 'styczniu'
        case '02':
            return 'lutym'
        case '03':
            return 'marcu'
        case '04': 
            return 'kwietniu'
        case '05': 
            return 'maju'    
        case '06': 
            return 'czerwcu'
        case '07': 
            return 'lipcu'
        case '08': 
            return 'sierpniu'
        case '09': 
            return 'wrześniu'
        case '10': 
            return 'październiku'
        case '11': 
            return 'listopadzie'
        case '12': 
            return 'grudniu'

doc = Document()

sections = doc.sections
for section in sections:
    section.left_margin = (Inches(1))
    section.right_margin = (Inches(1))

header = doc.sections[0].header
ht0=header.add_paragraph()
kh=ht0.add_run()
kh.add_picture('cgi.jpg')

footer = doc.sections[0].footer
ft0=footer.add_paragraph()
kf=ft0.add_run('Warsaw 2133430.2')
font = kf.font
font.name = 'Verdana'
font.size = Pt(8)

par1 = doc.add_paragraph()
run = par1.add_run('Załącznik Nr 1')
font = run.font
font.name = 'Times New Roman'
font.size = Pt(12)
font.bold = True

doc.add_paragraph()
doc.add_paragraph()

par2 = doc.add_paragraph()
paragraph_format = par2.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = par2.add_run('OŚWIADCZENIE PRACOWNIKA I ZGŁOSZENIE UTWORU')
font = run.font
font.name = 'Times New Roman'
font.size = Pt(12)
font.bold = True


doc.add_paragraph()

table = doc.add_table(rows=1, cols=1)
table.style = 'TableGrid'
cell = table.cell(0, 0)
cell.paragraphs[0].paragraph_format.space_before = Pt(4)
cell.paragraphs[0].text = "UWAGA:\n1.\tW formularzu pracownik wypełnia jedynie pola jasne.\n2.\tWszystkie pola jasne musza zostać wypełnione, aby utwór został zaakceptowany przez przełożonego oraz przez CGI.\n3.\tW jednym formularzu można zgłosić tylko jeden utwór. W przypadku gdy pracownik zgłasza więcej niż jeden utwór, dane kolejnych utworów należy podać w kolejnych formularzach, kopiując formularze i oznaczając każdy z nich kolejnym numerem porządkowym zgłaszanego utworu.\n4.\tW przypadku wskazania współautorów, pracownik zobowiązany jest do poinformowania współautorów o dokonanym zgłoszeniu utworu.\n5.\tFormularz powinien być zapisany na dysku publicznym „S” w formacie:\nImię i NaziwskoManagera_Oswiadczenie pracownika i zgloszenie utworu_Imię i Nazwisko Zgłaszającego Utwór.docx"
cell.paragraphs[0].paragraph_format.space_after = Pt(4)
font.name = 'Times New Roman'
font.size = Pt(10)


par3 = doc.add_paragraph()
par3 = doc.add_paragraph()
run = par3.add_run('Niniejszym przesyłam i zgłaszam następujące Utwory:')
font = run.font
font.name = 'Times New Roman'
font.size = Pt(12)

widths = [Inches(0.5), Inches(3), Inches(3)]

for i in range(number_of_works):
    par_next = doc.add_paragraph()
    par_next = doc.add_paragraph()
    run = par_next.add_run('Utwór Nr '+ str(i+1))
    font = run.font
    font.name = 'Times New Roman' 
    font.size = Pt(10)
    font.bold = True

    table = doc.add_table(rows=9, cols=3)
    table.style = 'TableGrid'
    table.autofit = False
    #table.allow_autofit = False

    cell = table.cell(0, 0)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "Lp."
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.cell(0,0).width = Inches(0.5)   

    cell = table.cell(0, 1)
    cell.paragraphs[0].text = ""
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.cell(0,1).width = Inches(3)   

    cell = table.cell(0, 2)
    cell.paragraphs[0].text = ""
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.cell(0,2).width = Inches(3.5)  

    cell = table.cell(1, 0)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "1."
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(1, 1)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "Imiona i nazwiska autora"
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[1].cells[1].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(1, 2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "Tomasz Skoczylas"
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    font.name = 'Times New Roman'
    font.size = Pt(10)

    cell = table.cell(2, 0)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "2."
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[2].cells[0].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(2, 1)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = 'Imiona i nazwiska współautorów'
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[2].cells[1].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(2, 2)

    cell = table.cell(3, 0)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "3."
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[3].cells[0].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(3, 1)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "Tytuł utworu (minimum pięć pierwszych słów z tytułu utworu)"
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[3].cells[1].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(3, 2)
    cell.width = Inches(3.5)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = mylist[i]
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    font.name = 'Times New Roman'
    font.size = Pt(10)

    cell = table.cell(4, 0)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "4."
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[4].cells[0].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(4, 1)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "Imię i nazwisko przełożonego odpowiedzialnego za wykonaną pracę"
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[4].cells[1].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(4, 2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "Paweł Sroczyński"
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    font.name = 'Times New Roman'
    font.size = Pt(10)

    cell = table.cell(5, 0)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "5."
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[5].cells[0].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(5, 1)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "Liczba godzin poświęconych na wykonanie utworu przez autora"
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[5].cells[1].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(5, 2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    # cell.paragraphs[0].text = str(particular_hours[i])
    cell.paragraphs[0].text = str(all_list[i])
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    font.name = 'Times New Roman'
    font.size = Pt(10)

    cell = table.cell(6, 0)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "6."
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[6].cells[0].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(6, 1)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "Opis utworu"
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[6].cells[1].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(6, 2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "Zestaw przypadków testowych weryfikujących transakcję TRANSACTION_NUMBER w systemie BiLaterals"
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    font.name = 'Times New Roman'
    font.size = Pt(10)

    cell = table.cell(7, 0)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text  = "7."
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[7].cells[0].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(7, 1)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "Link do dokumentu na dysku sieciowym lub wskazanie innej lokalizacji utworu"
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[7].cells[1].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(7, 2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "H:\\CW\\2022-" + period[:2]+ "\\" + mylist[i]
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    #cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].line_spacing_rule = WD_LINE_SPACING.DOUBLE

    font.name = 'Times New Roman'
    font.size = Pt(10)

    cell = table.cell(8, 0)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "8."
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[8].cells[0].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(8, 1)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = "Potwierdzenie powiadomienia  współautorów utworu o zgłoszeniu utworu"
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    font.name = 'Times New Roman'
    font.size = Pt(10)
    table.rows[8].cells[1].paragraphs[0].runs[0].font.bold = True

    cell = table.cell(8, 2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell.paragraphs[0].text = 'NIE'
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    font.name = 'Times New Roman'
    font.size = Pt(10)

    table_cells = table._cells
    for i in range(9):
        for j, cell in enumerate(table.rows[i].cells):
            cell.width = widths[j]
    for i in range(9):
        for j, cell in enumerate(table.rows[i].cells):
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            table.cell(i, 0)._tc.get_or_add_tcPr().append(shading_elm)
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            table.cell(i, 1)._tc.get_or_add_tcPr().append(shading_elm)
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm)



    # for row in table.rows:
    #     row.height = Cm(0.7)

# for i in range(9):
#     shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
#     table.cell(i, 0)._tc.get_or_add_tcPr().append(shading_elm)
#     table.cell(i, 1)._tc.get_or_add_tcPr().append(shading_elm)

doc.add_paragraph()

par4 = doc.add_paragraph()

run = par4.add_run('oraz potwierdzam, że jest/są to utwór/utwory chroniony/e na podstawie ustawy z dnia 4 lutego 1994 r. o prawie autorskim i prawach pokrewnych (tekst jednolity Dz.U. z 2000 r., Nr. 80, poz. 904 z późniejszymi zmianami).')
run.font.size = Pt(12)

par5 = doc.add_paragraph()

run = par5.add_run('Oświadczam, że praca została wykonana i ukończona w ostatnim miesiącu rozliczeniowym dla utworów chronionych prawami autorskimi tj. od 1. dnia bieżącego miesiąca do ostatniego dnia bieżącego miesiąca.')
run.font.size = Pt(12)

par6 = doc.add_paragraph()

run = par6.add_run('Ponadto, oświadczam, że w miesącu ')
run.font.size = Pt(12)

run = par6.add_run(w_miesiacu(period[:2]))
run.font.size = Pt(12)
run.font.bold = True

run = par6.add_run(' przepracowałem łącznie ')
run.font.size = Pt(12)

run = par6.add_run(str(period_working_hours))
run.font.size = Pt(12)
run.font.bold = True

run = par6.add_run(' godzin, z czego ')
run.font.size = Pt(12)

run = par6.add_run(str(int(period_working_hours*0.9)))
run.font.size = Pt(12)
run.font.bold = True

run = par6.add_run(' godzin poświęciłem na wykonanie wyżej wskazanych Utworów.')
run.font.size = Pt(12)

par7 = doc.add_paragraph()

run = par7.add_run('Tomasz Skoczylas')
run.font.size = Pt(12)

par8 = doc.add_paragraph()

run = par8.add_run('27.' + period[:2] + '.' + period[3:7])
run.font.size = Pt(12)
run.font.bold = True
run.font.italic = True

doc.save(cw_working_dir + '\P1aweł_Sroczyński_Oswiadczenie pracownika i Zgloszenie utworu_Tomasz_Skoczylas_' + period + '.docx')
